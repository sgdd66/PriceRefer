from docx import Document
from docx.shared import Inches

class CTDocx(object):
    def __init__(self,filePath):
        document=Document()
        self.filePath=filePath
        document.add_picture('icon1.png', width=Inches(2))
        document.add_heading('重庆通用（工业）集团有限责任公司', 0)
        document.add_paragraph('CHONGQING GENERAL INDUSTRY(GROUP) CO.,LTD.')
        self.document=document

    def ProgramTable(self,Data):

        table = self.document.add_table(rows=5, cols=2)
        cells = table.rows[0].cells
        cells[0].text = '项目名称'
        cells[1].text = Data[0]

        cells = table.rows[1].cells
        cells[0].text = '设计人'
        cells[1].text = Data[1]

        cells = table.rows[2].cells
        cells[0].text = '报告时间'
        cells[1].text = Data[2]

        cells = table.rows[3].cells
        cells[0].text = '风机型号'
        cells[1].text = Data[3]

        cells = table.rows[4].cells
        cells[0].text = '风机总重'
        cells[1].text = Data[4]
        self.document.add_paragraph()

    def ParamTable(self,Data):
        self.document.add_heading('风机使用参数', level=1)
        table = self.document.add_table(rows=4, cols=5)
        cells = table.rows[0].cells
        cells[0].text = '进气压力(Pa)'
        cells[1].text = '进气温度(K)'
        cells[2].text = '全压(K)'
        cells[3].text = '流量(m3/s)'
        cells[4].text = '转速(r/min)'

        cells = table.rows[1].cells
        cells[0].text = "{0}".format(Data[0])
        cells[1].text = "{0}".format(Data[1])
        cells[2].text = "{0}".format(Data[2])
        cells[3].text = "{0}".format(Data[3])
        cells[4].text = "{0}".format(Data[4])

        cells = table.rows[2].cells
        cells[0].text = '材料密度(kg/m3)'
        cells[1].text = '主轴+轴盘重量(kg)'
        cells[2].text = '其它重量(kg)'
        cells[3].text = '传动方式'
        cells[4].text = '叶轮直径(mm)'

        cells = table.rows[3].cells
        cells[0].text = "{0}".format(Data[5])
        cells[1].text = "{0}".format(Data[6])
        cells[2].text = "{0}".format(Data[7])
        cells[3].text = "{0}".format(Data[8])
        cells[4].text = "{0}".format(Data[9])

        self.document.add_paragraph()

    def ShellTable(self,Data):
        self.document.add_heading('机壳组', level=2)
        table = self.document.add_table(rows=5, cols=3)
        cells = table.rows[0].cells
        cells[1].text = '板厚(mm)'
        cells[2].text = '其它'

        cells = table.rows[1].cells
        cells[0].text = '进气箱'
        cells[1].text = "{0}".format(Data[0])
        cells[2].text = Data[1]

        cells = table.rows[2].cells
        cells[0].text = '蜗壳'
        cells[1].text = "{0}".format(Data[2])
        cells[2].text = Data[3]

        cells = table.rows[3].cells
        cells[0].text = '公用侧板'
        cells[1].text = "{0}".format(Data[4])
        cells[2].text = Data[5]

        cells = table.rows[4].cells
        cells[0].text = '后侧板'
        cells[1].text = "{0}".format(Data[6])
        cells[2].text = Data[7]
        self.document.add_paragraph(Data[8])

    def InletTable(self,Data):
        self.document.add_heading('进风口组', level=2)
        table = self.document.add_table(rows=3, cols=3)
        cells = table.rows[0].cells
        cells[1].text = '板厚(mm)'
        cells[2].text = '其它'

        cells = table.rows[1].cells
        cells[0].text = '法兰'
        cells[1].text = "{0}".format(Data[0])
        cells[2].text = Data[1]

        cells = table.rows[2].cells
        cells[0].text = '进风筒'
        cells[1].text = "{0}".format(Data[2])
        cells[2].text = Data[3]

        self.document.add_paragraph()

    def RotorTable(self, Data):
        self.document.add_heading('叶轮组', level=2)
        table = self.document.add_table(rows=4, cols=3)
        cells = table.rows[0].cells
        cells[1].text = '板厚(mm)'
        cells[2].text = '其它'

        cells = table.rows[1].cells
        cells[0].text = '前盘'
        cells[1].text = "{0}".format(Data[0])
        cells[2].text = Data[1]

        cells = table.rows[2].cells
        cells[0].text = '叶片'
        cells[1].text = "{0}".format(Data[2])
        cells[2].text = Data[3]

        cells = table.rows[3].cells
        cells[0].text = '中盘/后盘'
        cells[1].text = "{0}".format(Data[4])
        cells[2].text = Data[5]

        self.document.add_paragraph()

    def SaveFile(self):
        self.document.save(self.filePath)

    def sample(self):
        document = Document()

        document.add_heading('Document Title', 0)

        p = document.add_paragraph('A plain paragraph having some ')
        p.add_run('bold').bold = True
        p.add_run(' and some ')
        p.add_run('italic.').italic = True

        document.add_heading('Heading, level 1', level=1)
        document.add_paragraph('Intense quote', style='IntenseQuote')

        document.add_paragraph(
            'first item in unordered list', style='ListBullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='ListNumber'
        )

        document.add_picture('icon1.png', width=Inches(2.5))

        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for item in range(3):
            row_cells = table.add_row().cells
            row_cells[0].text = '11'
            row_cells[1].text = '22'
            row_cells[2].text = '33'

        document.add_page_break()

        document.save('F:/demo.docx')

if __name__=="__main__":
    test=CTDocx("F:/test.docx")
    Data=["项目名称","设计人","报告时间","风机型号","风机总重"]
    test.ProgramTable(Data)
    Data=["进气压力","进气温度","全压","流量","转速","材料密度","主轴+轴盘重量","其它重量","传动方式","叶轮直径"]
    test.ParamTable(Data)
    Data=["板厚","带法兰或加强筋","板厚","带法兰或加强筋","板厚","带法兰或加强筋","板厚","带法兰或加强筋","是否加脚板组"]
    test.ShellTable(Data)
    Data=["板厚","带法兰或加强筋","板厚","带法兰或加强筋"]
    test.InletTable(Data)
    Data=["板厚","带法兰或加强筋","板厚","带法兰或加强筋","板厚","带法兰或加强筋"]
    test.RotorTable(Data)
    test.SaveFile()